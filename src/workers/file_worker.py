"""File PII redaction worker — handles PDF and DOCX."""
from __future__ import annotations

import io
import logging
import time
from pathlib import Path
from typing import Any

import pdfplumber
import docx
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

from opentelemetry import trace

from src.observability.metrics import request_latency, worker_errors_total
from src.observability.tracer import get_tracer
from src.workers.text_worker import TextWorker

logger = logging.getLogger(__name__)
tracer = get_tracer("file_worker")


class FileWorker:
    """Redacts PII in PDF and DOCX files."""

    def __init__(self) -> None:
        self._text_worker = TextWorker()
        logger.info("FileWorker initialized")

    # ── PDF ─────────────────────────────────────────────────────────────────────

    def redact_pdf(
        self, pdf_bytes: bytes
    ) -> tuple[bytes, dict[str, str]]:
        """Extract text from PDF, redact PII, rebuild PDF. Returns (pdf_bytes, mapping)."""
        with tracer.start_as_current_span("file_worker.redact_pdf") as span:
            t0 = time.perf_counter()
            combined_mapping: dict[str, str] = {}
            pages_text: list[str] = []

            with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
                span.set_attribute("pdf.pages", len(pdf.pages))
                for page in pdf.pages:
                    raw_text = page.extract_text() or ""
                    redacted, mapping = self._text_worker.redact_safe(raw_text)
                    combined_mapping.update(mapping)
                    pages_text.append(redacted)

            rebuilt = self._build_pdf(pages_text)
            request_latency.labels(service="file_worker", content_type="pdf").observe(
                time.perf_counter() - t0
            )
            return rebuilt, combined_mapping

    def _build_pdf(self, pages: list[str]) -> bytes:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = []
        for page_text in pages:
            for line in page_text.splitlines():
                if line.strip():
                    story.append(Paragraph(line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"), styles["Normal"]))
        doc.build(story)
        return buf.getvalue()

    # ── DOCX ────────────────────────────────────────────────────────────────────

    def redact_docx(
        self, docx_bytes: bytes
    ) -> tuple[bytes, dict[str, str]]:
        """Redact PII in DOCX paragraphs in-place. Returns (docx_bytes, mapping)."""
        with tracer.start_as_current_span("file_worker.redact_docx") as span:
            t0 = time.perf_counter()
            combined_mapping: dict[str, str] = {}

            document = docx.api.Document(io.BytesIO(docx_bytes))
            span.set_attribute("docx.paragraphs", len(document.paragraphs))

            for para in document.paragraphs:
                if not para.text.strip():
                    continue
                redacted, mapping = self._text_worker.redact_safe(para.text)
                combined_mapping.update(mapping)
                if redacted != para.text:
                    self._replace_paragraph_text(para, redacted)

            for table in document.tables:
                for row in table.rows:
                    for cell in row.cells:
                        for para in cell.paragraphs:
                            redacted, mapping = self._text_worker.redact_safe(para.text)
                            combined_mapping.update(mapping)
                            if redacted != para.text:
                                self._replace_paragraph_text(para, redacted)

            out_buf = io.BytesIO()
            document.save(out_buf)
            request_latency.labels(service="file_worker", content_type="docx").observe(
                time.perf_counter() - t0
            )
            return out_buf.getvalue(), combined_mapping

    @staticmethod
    def _replace_paragraph_text(para: Any, new_text: str) -> None:
        """Replace paragraph text while preserving the first run's formatting."""
        for run in para.runs:
            run.text = ""
        if para.runs:
            para.runs[0].text = new_text
        else:
            para.add_run(new_text)

    def process(
        self,
        content: bytes,
        content_type: str,
    ) -> tuple[bytes, dict[str, str]]:
        """Dispatch to PDF or DOCX handler based on content_type."""
        try:
            if "pdf" in content_type:
                return self.redact_pdf(content)
            elif "docx" in content_type or "wordprocessingml" in content_type:
                return self.redact_docx(content)
            else:
                logger.warning("FileWorker: unsupported content_type=%s", content_type)
                return content, {}
        except Exception as exc:
            worker_errors_total.labels(worker="file", error_type=type(exc).__name__).inc()
            logger.error("FileWorker.process failed: content_type=%s err=%s", content_type, exc, exc_info=True)
            return content, {}
