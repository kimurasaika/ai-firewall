"""Re-pack DOCX/PDF files after de-anonymization."""
from __future__ import annotations

import io
import logging

import pdfplumber
import docx
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet
from reportlab.platypus import Paragraph, SimpleDocTemplate

from src.deanonymizer.deanonymizer import Deanonymizer

logger = logging.getLogger(__name__)


class FileRepacker:
    """Runs de-anonymization over file content and re-packs the file."""

    def __init__(self) -> None:
        self._deanon = Deanonymizer()

    def repack_pdf(
        self,
        pdf_bytes: bytes,
        mapping: dict[str, str],
    ) -> tuple[bytes, list[str]]:
        """Extract text from PDF → deanonymize → rebuild. Returns (bytes, misses)."""
        all_misses: list[str] = []
        pages_text: list[str] = []

        with pdfplumber.open(io.BytesIO(pdf_bytes)) as pdf:
            for page in pdf.pages:
                raw = page.extract_text() or ""
                restored, misses = self._deanon.deanonymize(raw, mapping)
                all_misses.extend(misses)
                pages_text.append(restored)

        rebuilt = self._build_pdf(pages_text)
        return rebuilt, all_misses

    def repack_docx(
        self,
        docx_bytes: bytes,
        mapping: dict[str, str],
    ) -> tuple[bytes, list[str]]:
        """Restore tokens in DOCX paragraphs in-place. Returns (bytes, misses)."""
        all_misses: list[str] = []
        document = docx.api.Document(io.BytesIO(docx_bytes))

        for para in document.paragraphs:
            restored, misses = self._deanon.deanonymize(para.text, mapping)
            all_misses.extend(misses)
            if restored != para.text:
                self._replace_para(para, restored)

        for table in document.tables:
            for row in table.rows:
                for cell in row.cells:
                    for para in cell.paragraphs:
                        restored, misses = self._deanon.deanonymize(para.text, mapping)
                        all_misses.extend(misses)
                        if restored != para.text:
                            self._replace_para(para, restored)

        out = io.BytesIO()
        document.save(out)
        return out.getvalue(), all_misses

    @staticmethod
    def _replace_para(para: object, new_text: str) -> None:
        for run in para.runs:  # type: ignore[attr-defined]
            run.text = ""
        if para.runs:  # type: ignore[attr-defined]
            para.runs[0].text = new_text  # type: ignore[attr-defined]
        else:
            para.add_run(new_text)  # type: ignore[attr-defined]

    @staticmethod
    def _build_pdf(pages: list[str]) -> bytes:
        buf = io.BytesIO()
        doc = SimpleDocTemplate(buf, pagesize=A4)
        styles = getSampleStyleSheet()
        story = [
            Paragraph(
                line.replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;"),
                styles["Normal"],
            )
            for page in pages
            for line in page.splitlines()
            if line.strip()
        ]
        doc.build(story)
        return buf.getvalue()

    def process(
        self,
        content: bytes,
        content_type: str,
        mapping: dict[str, str],
    ) -> tuple[bytes, list[str]]:
        ct = content_type.lower()
        if "pdf" in ct:
            return self.repack_pdf(content, mapping)
        if "docx" in ct or "wordprocessingml" in ct:
            return self.repack_docx(content, mapping)
        logger.warning("FileRepacker: unsupported content_type=%s", content_type)
        return content, []
